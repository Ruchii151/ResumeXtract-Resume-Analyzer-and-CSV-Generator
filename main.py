# import os
# import io
# import zipfile
# import pandas as pd
# import streamlit as st

# from dotenv import load_dotenv
# from typing import List

# from PyPDF2 import PdfReader
# from docx import Document
# from pydantic import BaseModel

# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.output_parsers import PydanticOutputParser

# # --------------------------------------------------
# # ENV
# # --------------------------------------------------
# load_dotenv()
# os.environ["GOOGLE_API_KEY"] = os.getenv("gemini")

# # --------------------------------------------------
# # SCHEMA (Pydantic)
# # --------------------------------------------------
# class ResumeSchema(BaseModel):
#     name: str
#     email: str
#     phone: str
#     skills: str
#     experience_summary: str
#     linkedin: str
#     github: str

# # --------------------------------------------------
# # OUTPUT PARSER
# # --------------------------------------------------
# parser = PydanticOutputParser(pydantic_object=ResumeSchema)

# # --------------------------------------------------
# # PROMPT (PLAIN STRING — NO TEMPLATES)
# # --------------------------------------------------
# def build_prompt(resume_text: str) -> str:
#     return f"""
# You are an AI resume analyzer.

# Extract information and return it STRICTLY in the following JSON format:
# {parser.get_format_instructions()}

# Rules:
# - Do NOT hallucinate
# - If information is missing, return an empty string
# - Output must strictly follow the schema
# - Return ONLY valid JSON, nothing else

# Resume:
# {resume_text}
# """

# # --------------------------------------------------
# # FILE READERS
# # --------------------------------------------------
# def read_pdf(file_bytes: bytes) -> str:
#     reader = PdfReader(io.BytesIO(file_bytes))
#     return "\n".join(page.extract_text() or "" for page in reader.pages)

# def read_docx(file_bytes: bytes) -> str:
#     doc = Document(io.BytesIO(file_bytes))
#     return "\n".join(p.text for p in doc.paragraphs)

# # --------------------------------------------------
# # ZIP EXTRACTION
# # --------------------------------------------------
# def extract_resumes(zip_file) -> List[dict]:
#     resumes = []

#     with zipfile.ZipFile(zip_file) as z:
#         for name in z.namelist():
#             if name.lower().endswith((".pdf", ".docx")):
#                 raw = z.read(name)
#                 text = read_pdf(raw) if name.endswith(".pdf") else read_docx(raw)

#                 resumes.append({
#                     "filename": name,
#                     "text": text
#                 })

#     return resumes

# # --------------------------------------------------
# # LLM PIPELINE
# # --------------------------------------------------
# def analyze_resume(resume_text: str) -> ResumeSchema:
#     llm = ChatGoogleGenerativeAI(
#         model="gemini-2.5-flash-lite",
#         temperature=0
#     )

#     prompt = build_prompt(resume_text)
#     response = llm.invoke(prompt)

#     return parser.parse(response.content)

# --------------------------------------------------
# STREAMLIT UI
# --------------------------------------------------
# st.set_page_config(
#     page_title="AI Resume Analyzer",
#     page_icon="📄",
#     layout="centered"
# )

# st.title("📄 AI-Powered Resume Analyzer")
# st.write("Upload a ZIP file containing multiple PDF/DOCX resumes.")

# uploaded_zip = st.file_uploader("Upload ZIP file", type=["zip"])

# if uploaded_zip:
#     with st.spinner("Analyzing resumes..."):
#         resumes = extract_resumes(uploaded_zip)
#         results = []

#         for resume in resumes:
#             parsed = analyze_resume(resume["text"])
#             row = parsed.model_dump()
#             row["filename"] = resume["filename"]
#             results.append(row)

#         df = pd.DataFrame(results)

#     st.success("Resume analysis completed ✅")
#     st.dataframe(df, use_container_width=True)

#     st.download_button(
#         label="⬇️ Download CSV",
#         data=df.to_csv(index=False),
#         file_name="resume_analysis.csv",
#         mime="text/csv"
#     )











# ----------------------------------------------- Above is simple version ---------------------------------------------------------------------------
# ----------------------------------------------- Below is enhanced version with better UI ----------------------------------------------------------------

import os
import io
import zipfile
import pandas as pd
import streamlit as st

from dotenv import load_dotenv
from typing import List

from PyPDF2 import PdfReader
from docx import Document
from pydantic import BaseModel

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.output_parsers import PydanticOutputParser

# --------------------------------------------------
# ENV
# --------------------------------------------------
load_dotenv()
os.environ["GOOGLE_API_KEY"] = os.getenv("gemini")

# --------------------------------------------------
# SCHEMA (Pydantic)
# --------------------------------------------------
class ResumeSchema(BaseModel):
    name: str
    email: str
    phone: str
    skills: str
    experience_summary: str
    linkedin: str
    github: str

# --------------------------------------------------
# OUTPUT PARSER
# --------------------------------------------------
parser = PydanticOutputParser(pydantic_object=ResumeSchema)

# --------------------------------------------------
# PROMPT (PLAIN STRING — NO TEMPLATES)
# --------------------------------------------------
def build_prompt(resume_text: str) -> str:
    return f"""
You are an AI resume analyzer.

Extract information and return it STRICTLY in the following JSON format:
{parser.get_format_instructions()}

Rules:
- Do NOT hallucinate
- If information is missing, return an empty string
- Output must strictly follow the schema
- Return ONLY valid JSON, nothing else

Resume:
{resume_text}
"""

# --------------------------------------------------
# FILE READERS
# --------------------------------------------------
def read_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def read_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)

# --------------------------------------------------
# ZIP EXTRACTION
# --------------------------------------------------
def extract_resumes(zip_file) -> List[dict]:
    resumes = []

    with zipfile.ZipFile(zip_file) as z:
        for name in z.namelist():
            if name.lower().endswith((".pdf", ".docx")):
                raw = z.read(name)
                text = read_pdf(raw) if name.endswith(".pdf") else read_docx(raw)

                resumes.append({
                    "filename": name,
                    "text": text
                })

    return resumes

# --------------------------------------------------
# LLM PIPELINE
# --------------------------------------------------
def analyze_resume(resume_text: str) -> ResumeSchema:
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash-lite",
        temperature=0
    )

    prompt = build_prompt(resume_text)
    response = llm.invoke(prompt)

    return parser.parse(response.content)

# --------------------------------------------------
# STREAMLIT UI (GRADIENT / NEON)
# --------------------------------------------------
st.set_page_config(
    page_title="AI Resume Analyzer",
    page_icon="📄",
    layout="wide"
)

# ---- Advanced Gradient CSS ----
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');

html, body, [class*="css"]  {
    font-family: 'Inter', sans-serif;
}

/* Background */
.main {
    background: radial-gradient(circle at top right, #7c2d12, #020617 60%);
}

/* Container */
.block-container {
    padding-top: 3rem;
    max-width: 1200px;
} 


/* Title */
.title {
    font-size: 3rem;
    font-weight: 800;
    background: linear-gradient(90deg, #fb7185, #facc15);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    margin-bottom: 0.3rem;
}

/* Subtitle */
.subtitle {
    color: #9ca3af;
    font-size: 1.1rem;
    margin-bottom: 2.5rem;
}

/* Button */
.stDownloadButton button {
    background: linear-gradient(90deg, #7c3aed, #06b6d4);
    color: white;
    border-radius: 12px;
    border: none;
    padding: 0.75rem 1.2rem;
    font-weight: 600;
    box-shadow: 0 0 20px rgba(124,58,237,0.6);
}

.stDownloadButton button:hover {
    transform: scale(1.03);
    box-shadow: 0 0 30px rgba(6,182,212,0.9);
}

/* File uploader */
section[data-testid="stFileUploader"] {
    border-radius: 14px;
    border: 1px dashed rgba(255,255,255,0.3);
    padding: 1.5rem;
}

/* Dataframe */
[data-testid="stDataFrame"] {
    border-radius: 14px;
    overflow: hidden;
}

/* Footer */
.footer {
    text-align: center;
    color: #6b7280;
    font-size: 0.85rem;
    margin-top: 3rem;
}
</style>
""", unsafe_allow_html=True)

# ---- Header ----
st.markdown('<div class="title">AI Resume Analyzer</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="subtitle">LLM-powered bulk resume extraction with structured outputs</div>',
    unsafe_allow_html=True
)

# ---- Upload Card ----
st.markdown('<div class="glass">', unsafe_allow_html=True)
st.markdown("### Upload Resume ZIP")
st.write("Drop a ZIP containing **PDF / DOCX resumes**. The AI handles the rest.")
uploaded_zip = st.file_uploader(
    "Upload ZIP",
    type=["zip"],
    label_visibility="collapsed"
)
st.markdown('</div>', unsafe_allow_html=True)

# ---- Processing ----
if uploaded_zip:
    with st.spinner("🧠Parsing resumes with Gemini…"):
        resumes = extract_resumes(uploaded_zip)
        results = []

        progress = st.progress(0)
        total = len(resumes)

        for i, resume in enumerate(resumes, start=1):
            parsed = analyze_resume(resume["text"])
            row = parsed.model_dump()
            row["filename"] = resume["filename"]
            results.append(row)
            progress.progress(i / total)

        df = pd.DataFrame(results)

    st.success(f"✨{len(df)} resumes analyzed successfully")

    # ---- Results Card ----
    st.markdown('<div class="glass">', unsafe_allow_html=True)
    st.markdown("### 📊Structured Resume Data")
    st.dataframe(df, use_container_width=True, hide_index=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ---- Download ----
    st.download_button(
        label="⬇️Download CSV",
        data=df.to_csv(index=False),
        file_name="resume_analysis.csv",
        mime="text/csv",
        use_container_width=True
    )

# ---- Footer ----
st.markdown(
    '<div class="footer">LangChain · Gemini · Streamlit · Pydantic</div>',
    unsafe_allow_html=True
)
